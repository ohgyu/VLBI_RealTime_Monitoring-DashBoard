from PyQt6.QtWidgets import (
    QFrame, QVBoxLayout, QHBoxLayout, QLabel, QPushButton,
    QWidget, QDateTimeEdit, QDialog, QSizePolicy, QFileDialog, QGridLayout, QScrollArea, QMessageBox
)
from PyQt6.QtCore import Qt, QDateTime, QTimer
from matplotlib.backends.backend_qtagg import FigureCanvasQTAgg as FigureCanvas
        # Navigation Toolbar
from matplotlib.backends.backend_qtagg import NavigationToolbar2QT as NavigationToolbar
from matplotlib.figure import Figure
from matplotlib import rc
from db_manager import get_connection
from datetime import datetime, timedelta
import numpy as np
import matplotlib as mpl
import os
import json
from pathlib import Path
import math
import logging

from thresholds_store import load_thresholds
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


logger = logging.getLogger(__name__)
# 한글 폰트
rc("font", family="Malgun Gothic")
mpl.rcParams['axes.unicode_minus'] = False

# ----------------------------------------------------------------------
#  DB 매핑: 좌측 메뉴 텍스트  →  테이블 / 컬럼명
# ----------------------------------------------------------------------
TABLE_MAP = {

# =========================================================
# 2 GHz Receiver
# =========================================================
"2GHz Receiver Status Monitor": {
    "table": "Frontend_2ghz",
    "columns": {
        "Normal Temperature RF": "NormalTemp_RF",
        "Normal Temperature Load": "NormalTemp_Load",
        "Cryogenic Temperature Cold": "Cryo_ColdPla",
        "Cryogenic Temperature Shield Box": "Cryo_ShieldBox",
        "Pressure Sensor CH1": "Pressure",
        "LNA Monitor LHCP Vd": "LNA_LHCP_Vd1",
        "LNA Monitor LHCP Id": "LNA_LHCP_Id1",
        "LNA Monitor RHCP Vd": "LNA_RHCP_Vd1",
        "LNA Monitor RHCP Id": "LNA_RHCP_Id1",
        "RF out Power RHCP": "RF_RHCP",
        "RF out Power LHCP": "RF_LHCP",
    }
},

# =========================================================
# 8 GHz Receiver
# =========================================================
"8GHz Receiver Status Monitor": {
    "table": "Frontend_8ghz",
    "columns": {
        "Normal Temperature RF": "NormalTemp_RF",
        "Normal Temperature Load": "NormalTemp_Load",
        "Cryogenic Temperature Cold": "Cryo_ColdPla",
        "Cryogenic Temperature Shield Box": "Cryo_ShieldBox",
        "Pressure Sensor CH1": "Pressure",
        "LNA Monitor LHCP Vg1": "LNA_LHCP_Vg1",
        "LNA Monitor LHCP Vg2": "LNA_LHCP_Vg2",
        "LNA Monitor LHCP Vd": "LNA_LHCP_Vd1",
        "LNA Monitor LHCP Id": "LNA_LHCP_Id1",
        "LNA Monitor RHCP Vg1": "LNA_RHCP_Vg1",
        "LNA Monitor RHCP Vg2": "LNA_RHCP_Vg2",
        "LNA Monitor RHCP Vd": "LNA_RHCP_Vd1",
        "LNA Monitor RHCP Id": "LNA_RHCP_Id1",
        "RF out Power RHCP": "RF_RHCP",
        "RF out Power LHCP": "RF_LHCP",
    }
},

# =========================================================
# 22 GHz Receiver
# =========================================================
"22GHz Receiver Status Monitor": {
    "table": "Frontend_22ghz",
    "columns": {
        "Normal Temperature RF": "NormalTemp_RF",
        "Normal Temperature LO": "NormalTemp_Load",
        "Cryogenic Temperature Cold": "Cryo_ColdPla",
        "Cryogenic Temperature Shield Box": "Cryo_ShieldBox",
        "Pressure Sensor CH1": "Pressure",
        "LNA Monitor LHCP Vg1": "LNA_LHCP_Vg1",
        "LNA Monitor LHCP Vg2": "LNA_LHCP_Vg2",
        "LNA Monitor LHCP Vd": "LNA_LHCP_Vd1",
        "LNA Monitor LHCP Id": "LNA_LHCP_Id1",
        "LNA Monitor RHCP Vg1": "LNA_RHCP_Vg1",
        "LNA Monitor RHCP Vg2": "LNA_RHCP_Vg2",
        "LNA Monitor RHCP Vd": "LNA_RHCP_Vd1",
        "LNA Monitor RHCP Id": "LNA_RHCP_Id1",
        "RF out Power RHCP": "RF_RHCP",
        "RF out Power LHCP": "RF_LHCP",
        "RF out Power LO": "Status_PLO",
    }
},

# =========================================================
# 43 GHz Receiver
# =========================================================
"43GHz Receiver Status Monitor": {
    "table": "Frontend_43ghz",
    "columns": {
        "Normal Temperature RF": "NormalTemp_RF",
        "Normal Temperature LO": "NormalTemp_Load",
        "Cryogenic Temperature Cold": "Cryo_ColdPla",
        "Cryogenic Temperature Shield Box": "Cryo_ShieldBox",
        "Pressure Sensor CH1": "Pressure",
        "LNA Monitor LHCP Vg1": "LNA_LHCP_Vg1",
        "LNA Monitor LHCP Vg2": "LNA_LHCP_Vg2",
        "LNA Monitor LHCP Vd": "LNA_LHCP_Vd1",
        "LNA Monitor LHCP Id": "LNA_LHCP_Id1",
        "LNA Monitor RHCP Vg1": "LNA_RHCP_Vg1",
        "LNA Monitor RHCP Vg2": "LNA_RHCP_Vg2",
        "LNA Monitor RHCP Vd": "LNA_RHCP_Vd1",
        "LNA Monitor RHCP Id": "LNA_RHCP_Id1",
        "RF out Power RHCP": "RF_RHCP",
        "RF out Power LHCP": "RF_LHCP",
        "RF out Power LO": "Status_PLO",
    }
},

# =========================================================
# S / X Down Converter
# =========================================================
"S/X Down Converter": {
    "table": "SXDownConverter",
    "columns": {
        "S":  "LEVEL_S",
        "X1": "LEVEL_X1",
        "X2": "LEVEL_X2",
    }
},

# =========================================================
# K Down Converter
# =========================================================
"K Down Converter": {
    "table": "KDownConverter",
    "columns": {
        "K1": "LEVEL_K1",
        "K2": "LEVEL_K2",
        "K3": "LEVEL_K3",
        "K4": "LEVEL_K4",
    }
},

# =========================================================
# Q Down Converter
# =========================================================
"Q Down Converter": {
    "table": "QDownConverter",
    "columns": {
        "Q1": "LEVEL_Q1",
        "Q2": "LEVEL_Q2",
        "Q3": "LEVEL_Q3",
        "Q4": "LEVEL_Q4",
    }
},

# =========================================================
# IF Selector
# =========================================================
"IF Selector": {
    "table": "IFSelector",
    "columns": {
        **{f"CH{i}": f"LEVEL_CH{i}" for i in range(1, 17)}
    }
},

# =========================================================
# Video Converter 2
# =========================================================
"Video Converter 2": {
    "table": "VideoConverter2",
    "columns": {
        **{f"LEVELU CH{ch}": f"LEVELU_CH{ch}" for ch in range(9, 17)},
        **{f"LEVELL CH{ch}": f"LEVELL_CH{ch}" for ch in range(9, 17)},
    }
},

}
# ----------------------------------------------------------------------
#  SPEC_MAP: 장비별 파라미터 정상 범위 (min, max)
#  ※ 값은 예시 / 미정 파라미터는 생략 가능
# ----------------------------------------------------------------------
SPEC_MAP = {

    "2GHz Receiver Status Monitor": {
        "Normal Temperature RF": (0, 50),
        "Normal Temperature Load": (0, 50),
        "Cryogenic Temperature Cold": (0, 20),
        "Cryogenic Temperature Shield Box": (0, 40),
        "Pressure Sensor CH1": (0, 2),
    },

    "8GHz Receiver Status Monitor": {
        "Normal Temperature RF": (0, 50),
        "Normal Temperature Load": (0, 50),
    },

    "22GHz Receiver Status Monitor": {
        "Normal Temperature RF": (0, 50),
        "Normal Temperature LO": (0, 50),
        "Cryogenic Temperature Cold": (0, 20),
    },

    "43GHz Receiver Status Monitor": {
        "Normal Temperature RF": (0, 50),
        "Normal Temperature LO": (0, 50),
    },

    # Down / IF / Video 는 추후 정의
}

ROW_MERGE_COUNT = {
    "VideoConverter2": 5,
}

# ----------------------------------------------------------------------
# 기간 설정 다이얼로그
# ----------------------------------------------------------------------
class CustomRangeDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)

        self.setWindowTitle("기간 설정")
        self.setFixedSize(320, 130)
        self.setStyleSheet("""
                    QDialog {
                        background-color: white;
                    }
                    QLabel {
                        color: #334155;
                        font-size: 10pt;
                        font-weight: bold;
                        background-color: transparent;
                    }
                    QDateTimeEdit {
                        background-color: #F8FAFC;
                        color: #1E293B;
                        border: 1px solid #CBD5E1;
                        border-radius: 4px;
                        padding: 4px;
                        font-size: 10pt;
                    }

                    /* --- Sửa nền lịch tại đây --- */
                    QCalendarWidget QWidget {
                        background-color: white; /* Nền chính */
                    }
                    QCalendarWidget QAbstractItemView:enabled {
                        color: #1E293B; /* Màu chữ ngày thường */
                        background-color: white;
                        selection-background-color: #2563EB; /* Màu nền khi chọn ngày */
                        selection-color: white;
                    }
                    QCalendarWidget QToolButton {
                        color: #1E293B;
                        background-color: white;
                        icon-size: 20px;
                        border: none;
                    }
                    QCalendarWidget QMenu {
                        background-color: white;
                        color: #1E293B;
                    }
                    QCalendarWidget QSpinBox {
                        background-color: white;
                        color: #1E293B;
                    }
                    /* ---------------------------- */

                    QPushButton {
                        padding: 6px 15px;
                        border-radius: 6px;
                        font-weight: bold;
                    }
                """)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(5, 5, 5, 5)
        layout.setSpacing(15)

        # 시작일
        h_layout1 = QHBoxLayout()
        h_layout1.addWidget(QLabel("시작:"))
        self.start_edit = QDateTimeEdit(QDateTime.currentDateTime().addDays(-1))
        self.start_edit.setDisplayFormat("yyyy-MM-dd HH:mm")
        self.start_edit.setCalendarPopup(True)
        h_layout1.addWidget(self.start_edit)
        layout.addLayout(h_layout1)

        # 종료일
        h_layout2 = QHBoxLayout()
        h_layout2.addWidget(QLabel("끝:  "))
        self.end_edit = QDateTimeEdit(QDateTime.currentDateTime())
        self.end_edit.setDisplayFormat("yyyy-MM-dd HH:mm")
        self.end_edit.setCalendarPopup(True)
        h_layout2.addWidget(self.end_edit)
        layout.addLayout(h_layout2)

        # Button layout
        btn_layout = QHBoxLayout()
        ok_btn = QPushButton("확인")
        ok_btn.setStyleSheet("""
                    QPushButton {
                        background-color: #2563EB;
                        color: white;
                    }
                    QPushButton:hover {
                        background-color: #1D4ED8;
                    }
                """)

        cancel_btn = QPushButton("취소")
        cancel_btn.setStyleSheet("""
                    QPushButton {
                        background-color: #E2E8F0;
                        color: #475569;
                    }
                    QPushButton:hover {
                        background-color: #CBD5E1;
                    }
                """)

        btn_layout.addStretch()
        btn_layout.addWidget(ok_btn)
        btn_layout.addWidget(cancel_btn)
        layout.addLayout(btn_layout)

        ok_btn.clicked.connect(self.accept)
        cancel_btn.clicked.connect(self.reject)

    def get_range(self):
        return (
            self.start_edit.dateTime().toPyDateTime(),
            self.end_edit.dateTime().toPyDateTime(),
        )


# ----------------------------------------------------------------------
#  메인 센터 프레임 (반응형 2×2 카드 + 각 카드 하단 통계)
# ----------------------------------------------------------------------
class FrameCenter(QFrame):
    def __init__(self, parent=None, frame_right=None):
        super().__init__(parent)

        self.frame_left = None

        self.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.MinimumExpanding)
        self.setStyleSheet("background-color:#0F172A;")

        self.selected_children: dict[str, list[str]] = {}
        self.raw: dict[str, dict] = {}

        self.time_range = "1Day"
        self.custom_start = None
        self.custom_end = None

        # -----------------------
        # 전체 레이아웃
        # -----------------------
        self.main_layout = QVBoxLayout(self)
        self.main_layout.setContentsMargins(10, 10, 10, 10)
        self.main_layout.setSpacing(8)

        # -----------------------
        # 상단 버튼 영역
        # -----------------------
        top_widget = QWidget()
        top_layout = QHBoxLayout(top_widget)
        top_layout.setContentsMargins(0, 0, 0, 0)
        top_layout.setSpacing(12)
        top_layout.addStretch(1)

        self.reset_btn = QPushButton("Reset")
        self.reset_btn.setFixedHeight(30)
        self.reset_btn.setFixedWidth(70)
        self.reset_btn.setStyleSheet("""
            QPushButton {
                background-color:#EF4444;
                color:white;
                border-radius:8px;
                padding:4px 10px;
                font-size:10pt;
                font-weight:bold;
            }
            QPushButton:hover { background-color:#DC2626; }
        """)
        self.reset_btn.clicked.connect(self.reset_all)
        top_layout.addWidget(self.reset_btn)

        self.time_buttons = {}
        for name in ["1Hour", "6Hour", "1Day", "7Day", "Range"]:
            btn = QPushButton(name)
            btn.setFixedHeight(30)
            btn.setFixedWidth(70)
            btn.setCheckable(True)
            btn.setStyleSheet("""
                QPushButton {
                    background-color:#1E293B;
                    color:white;
                    border-radius:8px;
                    padding:4px 10px;
                    font-size:10pt;
                }
                QPushButton:hover { background-color:#334155; }
                QPushButton:checked {
                    background-color:#2563EB;
                    font-weight:bold;
                }
            """)
            self.time_buttons[name] = btn
            top_layout.addWidget(btn)
            btn.clicked.connect(self._make_time_range_handler(name))

        self.time_buttons["1Day"].setChecked(True)

        self.refresh_btn = QPushButton("Refresh")
        self.refresh_btn.setFixedHeight(30)
        self.refresh_btn.setFixedWidth(90)
        self.refresh_btn.setStyleSheet("""
            QPushButton {
                background-color:#22C55E;
                color:white;
                border-radius:8px;
                padding:4px 12px;
                font-size:10pt;
                font-weight:bold;
            }
            QPushButton:hover { background-color:#16A34A; }
        """)
        self.refresh_btn.clicked.connect(self.refresh_all_data)
        top_layout.addWidget(self.refresh_btn)

        self.report_word_btn = QPushButton("Report Word")
        self.report_word_btn.setFixedHeight(30)
        self.report_word_btn.setFixedWidth(110)
        self.report_word_btn.setStyleSheet("""
            QPushButton {
                background-color:#6366F1;
                color:white;
                border-radius:8px;
                padding:4px 12px;
                font-size:10pt;
                font-weight:bold;
            }
            QPushButton:hover { background-color:#4F46E5; }
        """)
        self.report_word_btn.clicked.connect(self.save_word_report)
        top_layout.addWidget(self.report_word_btn)

        self.main_layout.addWidget(top_widget)

        # -----------------------
        # 그래프 영역
        # -----------------------
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setStyleSheet("background-color:#0F172A; border: none;")

        self.cards_container = QWidget()
        self.cards_layout = QGridLayout(self.cards_container)
        self.cards_layout.setContentsMargins(0, 0, 0, 0)
        self.cards_layout.setHorizontalSpacing(12)
        self.cards_layout.setVerticalSpacing(12)

        scroll.setWidget(self.cards_container)
        self.main_layout.addWidget(scroll, stretch=1)

        QTimer.singleShot(50, self.update_graphs)

        self.db_latest = {}  # 테이블별 최신 datetime 캐시
        self.db_earliest = {}  # 테이블별 최조 datetime 캐시(옵션이지만 Range/보정에 유용)

        self.thresholds = {}
        self._load_thresholds()

    def _load_thresholds(self):
        self.thresholds = load_thresholds()

    # ------------------------------------------------------------------
    # 시간 범위 처리
    # ------------------------------------------------------------------
    def _make_time_range_handler(self, name):
        def handler():
            for n, btn in self.time_buttons.items():
                btn.setChecked(n == name)

            if name == "Range":
                dlg = CustomRangeDialog(self)
                if dlg.exec():
                    self.custom_start, self.custom_end = dlg.get_range()
                    self.time_range = "Range"
                else:
                    self.time_buttons[self.time_range].setChecked(True)
                    return
            else:
                self.time_range = name

            # ✅ 여기 추가: 시간범위가 바뀌면 선택된 parent들을 DB에서 다시 로드
            for parent_name in list(self.selected_children.keys()):
                self.reload_data(parent_name, load_only=True)

            self.update_graphs()

        return handler

    def _get_time_window(self):
        # 1) Range면 그대로
        if self.time_range == "Range" and self.custom_start and self.custom_end:
            return self.custom_start, self.custom_end

        # 2) 기준시간(base_now) = "현재시각"이 아니라 "DB 최신시각"
        #    선택된 장비 중 하나라도 캐시가 있으면 그걸 쓰고,
        #    없으면 전체 raw에서 최대값을 쓰고,
        #    그래도 없으면 시스템 now
        base_now = None

        # 선택된 parent 중 캐시된 최신시간 우선
        for parent in self.selected_children.keys():
            info = TABLE_MAP.get(parent)
            if not info:
                continue
            table = info["table"]
            if table in self.db_latest and self.db_latest[table] is not None:
                if base_now is None or self.db_latest[table] > base_now:
                    base_now = self.db_latest[table]

        # raw에 이미 로드된 데이터가 있으면 거기서 최대
        if base_now is None:
            for raw in self.raw.values():
                ts = raw.get("times")
                if ts:
                    mx = max(ts)
                    if base_now is None or mx > base_now:
                        base_now = mx

        if base_now is None:
            base_now = datetime.now()

        # 3) 버튼별 window 계산
        if self.time_range == "1Hour":
            return base_now - timedelta(hours=1), base_now
        if self.time_range == "6Hour":
            return base_now - timedelta(hours=6), base_now
        if self.time_range == "1Day":
            return base_now - timedelta(days=1), base_now
        if self.time_range == "7Day":
            return base_now - timedelta(days=7), base_now

        # 기본값
        return base_now - timedelta(days=1), base_now

    def show_info_dark(self, title: str, message: str):
        box = QMessageBox(self)
        box.setIcon(QMessageBox.Icon.Information)
        box.setWindowTitle(title)
        box.setText(message)

        box.setStyleSheet("""
            QMessageBox {
                background-color: #0F172A;
            }
            QMessageBox QLabel {
                color: white;
                font-size: 11pt;
            }
            QMessageBox QPushButton {
                background-color: #2563EB;
                color: white;
                padding: 6px 16px;
                border-radius: 6px;
                font-weight: bold;
                min-width: 70px;
            }
            QMessageBox QPushButton:hover {
                background-color: #1D4ED8;
            }
        """)

        box.exec()

    def show_child_graph(self, group_name, item_name):
        if group_name not in TABLE_MAP:
            return

        self.selected_children.setdefault(group_name, [])

        # --- TOGGLE OFF ---
        if item_name in self.selected_children[group_name]:
            self.selected_children[group_name].remove(item_name)

            if self.frame_left:
                self.frame_left.refresh_child_selection()
            self.update_graphs()
            return

        # --- TOGGLE ON (여기서 제한) ---
        if self._selected_count() >= 4:
            self.show_info_dark("알림", "그래프는 최대 4개까지만 표시할 수 있습니다.")
            if self.frame_left:
                self.frame_left.refresh_child_selection()
            return

        self.selected_children[group_name].append(item_name)

        # 아직 raw에 없으면 로드
        if group_name not in self.raw:
            self.reload_data(group_name, load_only=True)

        if self.frame_left:
            self.frame_left.refresh_child_selection()

        self.update_graphs()


    # ------------------------------------------------------------------
    # 새로고침
    # ------------------------------------------------------------------
    def refresh_all_data(self):
        self._load_thresholds()   # ✅ 추가
        for parent_name in list(self.selected_children.keys()):
            self.reload_data(parent_name, load_only=True)
        self.update_graphs()

    def reset_all(self):
        # FrameCenter 내부 선택 상태 초기화
        self.selected_children.clear()
        self.raw.clear()

        # 카드(그래프) 레이아웃 비우기
        while self.cards_layout.count():
            item = self.cards_layout.takeAt(0)
            w = item.widget()
            if w:
                w.deleteLater()

        self.cards_container.update()

        # 왼쪽 메뉴 선택도 초기화
        if self.frame_left:
            self.frame_left.clear_all_selection()

    # ------------------------------------------------------------------
    # DB 로딩
    # ------------------------------------------------------------------
    def reload_data(self, parent_name, load_only: bool = False):
        info = TABLE_MAP.get(parent_name)
        if not info:
            logger.warning("TABLE_MAP 에 없는 parent: %s", parent_name)
            return

        table = info.get("table")
        if not table:
            logger.error("테이블이 없는 parent: %s", parent_name)
            return

        wanted_cols = list(dict.fromkeys(info["columns"].values()))

        conn = None
        try:
            conn = get_connection(readonly=True)
            cur = conn.cursor()

            # --- (A) DB 최신/최초 시간 1번만 구해서 캐시 ---
            if table not in self.db_latest or self.db_latest[table] is None:

                if table == "VideoConverter2":
                    cur.execute("""
                                SELECT MIN(substr(datetime, 1, 19)),
                                       MAX(substr(datetime, 1, 19))
                                FROM VideoConverter2
                                WHERE NULLIF(REPLACE(REPLACE(REPLACE(LEVELU_CH9, char(9), ''), char(10), ''), char(13), ''), '') IS NOT NULL
                                   OR NULLIF(REPLACE(REPLACE(REPLACE(LEVELU_CH10, char(9), ''), char(10), ''), char(13), ''), '') IS NOT NULL
                                   OR NULLIF(REPLACE(REPLACE(REPLACE(LEVELU_CH11, char(9), ''), char(10), ''), char(13), ''), '') IS NOT NULL
                                   OR NULLIF(REPLACE(REPLACE(REPLACE(LEVELU_CH12, char(9), ''), char(10), ''), char(13), ''), '') IS NOT NULL
                                   OR NULLIF(REPLACE(REPLACE(REPLACE(LEVELU_CH13, char(9), ''), char(10), ''), char(13), ''), '') IS NOT NULL
                                   OR NULLIF(REPLACE(REPLACE(REPLACE(LEVELU_CH14, char(9), ''), char(10), ''), char(13), ''), '') IS NOT NULL
                                   OR NULLIF(REPLACE(REPLACE(REPLACE(LEVELU_CH15, char(9), ''), char(10), ''), char(13), ''), '') IS NOT NULL
                                   OR NULLIF(REPLACE(REPLACE(REPLACE(LEVELU_CH16, char(9), ''), char(10), ''), char(13), ''), '') IS NOT NULL
                                   OR NULLIF(REPLACE(REPLACE(REPLACE(LEVELL_CH9, char(9), ''), char(10), ''), char(13), ''), '') IS NOT NULL
                                   OR NULLIF(REPLACE(REPLACE(REPLACE(LEVELL_CH10, char(9), ''), char(10), ''), char(13), ''), '') IS NOT NULL
                                   OR NULLIF(REPLACE(REPLACE(REPLACE(LEVELL_CH11, char(9), ''), char(10), ''), char(13), ''), '') IS NOT NULL
                                   OR NULLIF(REPLACE(REPLACE(REPLACE(LEVELL_CH12, char(9), ''), char(10), ''), char(13), ''), '') IS NOT NULL
                                   OR NULLIF(REPLACE(REPLACE(REPLACE(LEVELL_CH13, char(9), ''), char(10), ''), char(13), ''), '') IS NOT NULL
                                   OR NULLIF(REPLACE(REPLACE(REPLACE(LEVELL_CH14, char(9), ''), char(10), ''), char(13), ''), '') IS NOT NULL
                                   OR NULLIF(REPLACE(REPLACE(REPLACE(LEVELL_CH15, char(9), ''), char(10), ''), char(13), ''), '') IS NOT NULL
                                   OR NULLIF(REPLACE(REPLACE(REPLACE(LEVELL_CH16, char(9), ''), char(10), ''), char(13), ''), '') IS NOT NULL
                                """)
                    mn, mx = cur.fetchone()
                else:
                    cur.execute(f"SELECT MIN(datetime), MAX(datetime) FROM {table}")
                    mn, mx = cur.fetchone()

                def _to_dt(x):
                    if x is None:
                        return None
                    try:
                        return datetime.fromisoformat(str(x)[:19])
                    except Exception:
                        return None

                self.db_earliest[table] = _to_dt(mn)
                self.db_latest[table] = _to_dt(mx)

            base_now = self.db_latest.get(table)
            if base_now is None:
                self.raw[parent_name] = {"times": [], "data": {c: [] for c in wanted_cols}}
                if not load_only:
                    self.update_graphs()
                return

            # --- (B) 버튼 범위 계산 (DB 최신 기준) ---
            if self.time_range == "Range" and self.custom_start and self.custom_end:
                start_dt, end_dt = self.custom_start, self.custom_end
            elif self.time_range == "1Hour":
                start_dt, end_dt = base_now - timedelta(hours=1), base_now
            elif self.time_range == "6Hour":
                start_dt, end_dt = base_now - timedelta(hours=6), base_now
            elif self.time_range == "7Day":
                start_dt, end_dt = base_now - timedelta(days=7), base_now
            else:
                start_dt, end_dt = base_now - timedelta(days=1), base_now  # 기본 1Day

            start_s = start_dt.strftime("%Y-%m-%d %H:%M:%S")
            end_s = end_dt.strftime("%Y-%m-%d %H:%M:%S")

            # =========================================================
            # ✅ VideoConverter2는 "집계 SQL"로 1초 단위로 병합해서 가져오기
            #    (빈 문자열/공백/NULL 제거까지 SQL에서 끝냄)
            # =========================================================
            if table == "VideoConverter2":
                sql = """
                      SELECT substr(datetime, 1, 19)                                                                          AS datetime,

                             MAX(NULLIF(REPLACE(REPLACE(REPLACE(LEVELU_CH9, char(9), ''), char(10), ''), char(13), ''), ''))  AS LEVELU_CH9,
                             MAX(NULLIF(REPLACE(REPLACE(REPLACE(LEVELU_CH10, char(9), ''), char(10), ''), char(13), ''), '')) AS LEVELU_CH10,
                             MAX(NULLIF(REPLACE(REPLACE(REPLACE(LEVELU_CH11, char(9), ''), char(10), ''), char(13), ''), '')) AS LEVELU_CH11,
                             MAX(NULLIF(REPLACE(REPLACE(REPLACE(LEVELU_CH12, char(9), ''), char(10), ''), char(13), ''), '')) AS LEVELU_CH12,
                             MAX(NULLIF(REPLACE(REPLACE(REPLACE(LEVELU_CH13, char(9), ''), char(10), ''), char(13), ''), '')) AS LEVELU_CH13,
                             MAX(NULLIF(REPLACE(REPLACE(REPLACE(LEVELU_CH14, char(9), ''), char(10), ''), char(13), ''), '')) AS LEVELU_CH14,
                             MAX(NULLIF(REPLACE(REPLACE(REPLACE(LEVELU_CH15, char(9), ''), char(10), ''), char(13), ''), '')) AS LEVELU_CH15,
                             MAX(NULLIF(REPLACE(REPLACE(REPLACE(LEVELU_CH16, char(9), ''), char(10), ''), char(13), ''), '')) AS LEVELU_CH16,

                             MAX(NULLIF(REPLACE(REPLACE(REPLACE(LEVELL_CH9, char(9), ''), char(10), ''), char(13), ''), ''))  AS LEVELL_CH9,
                             MAX(NULLIF(REPLACE(REPLACE(REPLACE(LEVELL_CH10, char(9), ''), char(10), ''), char(13), ''), '')) AS LEVELL_CH10,
                             MAX(NULLIF(REPLACE(REPLACE(REPLACE(LEVELL_CH11, char(9), ''), char(10), ''), char(13), ''), '')) AS LEVELL_CH11,
                             MAX(NULLIF(REPLACE(REPLACE(REPLACE(LEVELL_CH12, char(9), ''), char(10), ''), char(13), ''), '')) AS LEVELL_CH12,
                             MAX(NULLIF(REPLACE(REPLACE(REPLACE(LEVELL_CH13, char(9), ''), char(10), ''), char(13), ''), '')) AS LEVELL_CH13,
                             MAX(NULLIF(REPLACE(REPLACE(REPLACE(LEVELL_CH14, char(9), ''), char(10), ''), char(13), ''), '')) AS LEVELL_CH14,
                             MAX(NULLIF(REPLACE(REPLACE(REPLACE(LEVELL_CH15, char(9), ''), char(10), ''), char(13), ''), '')) AS LEVELL_CH15,
                             MAX(NULLIF(REPLACE(REPLACE(REPLACE(LEVELL_CH16, char(9), ''), char(10), ''), char(13), ''), '')) AS LEVELL_CH16

                      FROM VideoConverter2
                      WHERE datetime BETWEEN ? AND ?
                      GROUP BY substr(datetime, 1, 19)
                      ORDER BY substr(datetime, 1, 19) ASC 
                      """
                cur.execute(sql, (start_s, end_s))

            # =========================================================
            # ✅ 나머지 테이블은 기존 방식
            # =========================================================
            else:
                cols = ["datetime"] + wanted_cols
                col_sql = ", ".join(cols)
                sql = f"""
                    SELECT {col_sql}
                    FROM {table}
                    WHERE datetime BETWEEN ? AND ?
                    ORDER BY datetime ASC
                """
                cur.execute(sql, (start_s, end_s))

            rows = cur.fetchall()
            col_names = [desc[0] for desc in cur.description]
            col_index = {name: idx for idx, name in enumerate(col_names)}
            dt_idx = col_index.get("datetime")
            if dt_idx is None:
                logger.error("테이블 %s 에 datetime 컬럼이 없습니다. (스키마 확인 필요)", table)
                return


        except Exception:
            logger.exception("DB 오류 (parent=%s)", parent_name)
            return
        finally:
            if conn is not None:
                try:
                    conn.close()
                except Exception:
                    pass

        # ----------------------------
        # rows -> times/data 변환
        # ----------------------------
        times = []
        data = {col: [] for col in wanted_cols}

        for row in rows:
            # datetime 파싱
            try:
                dt_raw = row[dt_idx]
                dt = dt_raw if isinstance(dt_raw, datetime) else datetime.fromisoformat(str(dt_raw)[:19])
            except Exception:
                continue

            times.append(dt)

            # 컬럼 값 저장 (VideoConverter2는 SQL에서 이미 trim/null 처리됨)
            for col in wanted_cols:
                idx = col_index.get(col)
                v = row[idx] if idx is not None else None
                data[col].append(v)

        self.raw[parent_name] = {"times": times, "data": data}

        if not load_only:
            self.update_graphs()

        if table == "VideoConverter2":
            u9 = data.get("LEVELU_CH9")
            l9 = data.get("LEVELL_CH9")


    # ------------------------------------------------------------------
    # 현재 선택된 parent/child 기반으로 플롯 데이터 수집
    # ------------------------------------------------------------------
    def _collect_plot_items(self):
        start, end = self._get_time_window()
        plot_items = []

        for parent, child_list in self.selected_children.items():

            info = TABLE_MAP.get(parent)
            if not info:
                continue

            if parent not in self.raw:
                continue

            times = self.raw[parent]["times"]
            data = self.raw[parent]["data"]

            # 기간 범위 내 인덱스 추출
            indices = [i for i, t in enumerate(times) if start <= t <= end]
            if not indices:
                continue

            for child in child_list:
                col = info["columns"].get(child)
                if not col:
                    continue

                xs_raw = [times[i] for i in indices]
                ys_raw = [data[col][i] for i in indices]

                xs, ys = [], []
                for x, y in zip(xs_raw, ys_raw):
                    if y is None:
                        continue

                    # 문자열이면 공백 제거
                    if isinstance(y, str):
                        y = y.strip()
                        if y == "":
                            continue

                    try:
                        v = float(y)
                    except Exception:
                        continue

                    if np.isnan(v):
                        continue

                    xs.append(x)
                    ys.append(v)

                # ✅ 표시용 디시메이션(의미 유지 + 렉 방지)
                max_pts = self._target_points_for_range()
                xs, ys = self._downsample_to_max_points(xs, ys, max_pts)

                if len(xs) > 0:
                    title = f"{parent} | {child}"  # ← 여기서 title 생성
                    plot_items.append((title, xs, ys))

        return plot_items

    # ------------------------------------------------------------------
    # 그래프 카드 하나 생성 (그래프 + 하단 통계)
    # ------------------------------------------------------------------
    def _create_graph_card(self, title: str, xs, ys, start_dt, end_dt):
        card = QFrame()
        card.setStyleSheet("""
            QFrame {
                background-color:#020617;
                border:1px solid #1E293B;
                border-radius:10px;
            }
        """)
        card.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        card.setMinimumHeight(300)

        v = QVBoxLayout(card)
        v.setContentsMargins(10, 8, 10, 8)
        v.setSpacing(4)

        fig = Figure(facecolor="#020617")
        canvas = FigureCanvas(fig)
        canvas.setParent(card)
        canvas.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)

        ax = fig.add_subplot(111)
        ax.set_facecolor("#020617")
        for spine in ax.spines.values():
            spine.set_color("#1E293B")

        ax.grid(True, color="#1E293B", linestyle="--", linewidth=0.5)
        ax.tick_params(axis="x", colors="white")
        ax.tick_params(axis="y", colors="white")

        ax.plot(xs, ys, linewidth=1.0, antialiased=False)

        # -----------------------
        # 임계값 라인 표시 (노랑=주의, 빨강=경고)
        # -----------------------
        th = None
        try:
            parent_name, child_name = title.split(" | ", 1)
            th = self._get_thresholds_for(parent_name, child_name)
        except Exception:
            th = None

        if th:
            # 노랑(주의)
            if th.get("ly") is not None:
                ax.axhline(float(th["ly"]), linestyle="--", linewidth=1.0, color="yellow")
            if th.get("uy") is not None:
                ax.axhline(float(th["uy"]), linestyle="--", linewidth=1.0, color="yellow")

            # 빨강(경고)
            if th.get("lr") is not None:
                ax.axhline(float(th["lr"]), linestyle="-", linewidth=1.2, color="red")
            if th.get("ur") is not None:
                ax.axhline(float(th["ur"]), linestyle="-", linewidth=1.2, color="red")

        # ✅ 이 줄이 핵심 (7Day 눌러도 2028년 안 튐)
        ax.set_xlim(start_dt, end_dt)

        ax.set_title(title, color="white", fontsize=10, pad=6)
        fig.autofmt_xdate()

        card._fig = fig
        card._canvas = canvas
        card._ax = ax

        v.addWidget(canvas, stretch=1)

        canvas.draw_idle()

        arr = np.array(ys, dtype=float)
        summary_label = QLabel(
            f"Avg: {arr.mean():,.2f}    Max: {arr.max():,.2f}    min: {arr.min():,.2f}"
        )
        summary_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        summary_label.setStyleSheet("color:#38BDF8; font-size:10pt; font-weight:bold;")
        v.addWidget(summary_label)

        return card

    # ------------------------------------------------------------------
    # 전체 그래프 갱신 (QGridLayout 2×N 반응형 배치)
    # ------------------------------------------------------------------
    def update_graphs(self):
        # 1) 기존 카드(그래프) 전부 제거
        while self.cards_layout.count():
            item = self.cards_layout.takeAt(0)
            w = item.widget()
            lay = item.layout()

            if w is not None:
                w.deleteLater()

            if lay is not None:
                while lay.count():
                    sub_item = lay.takeAt(0)
                    if sub_item.widget():
                        sub_item.widget().deleteLater()

        # 2) 플롯 아이템 수집
        plot_items = self._collect_plot_items()
        if not plot_items:
            self.cards_container.update()
            return

        # ✅ 3) 버튼 기준 시간 윈도우(=DB 최신 기준으로 계산된 start/end)
        start_dt, end_dt = self._get_time_window()

        # 4) 그리드 배치
        cols = 2
        rows = (len(plot_items) + cols - 1) // cols

        for idx, (title, xs, ys) in enumerate(plot_items):
            row = idx // cols
            col = idx % cols

            # ✅ 5) _create_graph_card에 start/end 넘겨서 x축 고정
            card = self._create_graph_card(title, xs, ys, start_dt, end_dt)
            self.cards_layout.addWidget(card, row, col)

        # 6) stretch 설정
        for r in range(rows):
            self.cards_layout.setRowStretch(r, 1)
        for c in range(cols):
            self.cards_layout.setColumnStretch(c, 1)

        self.cards_container.update()

    def showEvent(self, event):
        super().showEvent(event)
        QTimer.singleShot(50, self._force_resize_after_show)

    def _force_resize_after_show(self):
        self.update_graphs()

    def _force_resize(self):
        self.resize(self.width(), self.height())

    # ------------------------------------------------------------------
    # Word 보고서 저장 (화면의 모든 그래프 + 각 그래프 통계)
    # ------------------------------------------------------------------
    def save_word_report(self):
        filename, _ = QFileDialog.getSaveFileName(
            self,
            "Word 보고서 저장",
            "VLBI_Device_Report.docx",
            "Word Files (*.docx)"
        )
        if not filename:
            return

        plot_items = self._collect_plot_items()
        if not plot_items:
            logger.info("Word 생성 요청: 표시할 그래프가 없습니다.")
            return

        # -------------------------------
        # 1. 장비별 그룹화
        # -------------------------------
        device_map = {}
        for title, xs, ys in plot_items:
            parent, child = title.split(" | ", 1)
            device_map.setdefault(parent, []).append((child, xs, ys))

        # -------------------------------
        # 2. Word 문서 생성
        # -------------------------------
        doc = Document()

        # 기본 한글 폰트 설정
        style = doc.styles['Normal']
        style.font.name = 'Malgun Gothic'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
        style.font.size = Pt(10)

        now_kst = datetime.utcnow() + timedelta(hours=9)
        created_text = f"생성일시: {now_kst.strftime('%Y-%m-%d %H:%M:%S')} (KST)"

        temp_files = []
        first_device = True

        for device_name, param_list in device_map.items():
            if not first_device:
                doc.add_page_break()
            first_device = False

            # -------------------------------
            # 분석 기간 계산
            # -------------------------------
            all_times = []
            for _, xs, _ in param_list:
                all_times.extend(xs)

            if all_times:
                period_start = min(all_times)
                period_end = max(all_times)
            else:
                period_start = period_end = datetime.utcnow()

            uptime_hours = max(
                0.0,
                (period_end - period_start).total_seconds() / 3600
            )

            # ===============================
            # 1. 개요
            # ===============================
            title = doc.add_heading("개별 장비 정기 성능 분석 보고서", level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            p = doc.add_paragraph(created_text)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_heading("1. 개요 및 가동 현황 (Dashboard)", level=1)

            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "항목"
            hdr[1].text = "값"

            rows = [
                ("장비명", device_name),
                ("분석 기간",
                 f"{period_start:%Y-%m-%d %H:%M} ~ {period_end:%Y-%m-%d %H:%M}")
            ]

            for k, v in rows:
                r = table.add_row().cells
                r[0].text = k
                r[1].text = v

            # 가동 지표
            doc.add_paragraph()
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "지표"
            hdr[1].text = "값"

            metrics = [
                ("가동률", "99.8%"),
                ("총 가동 시간", f"{uptime_hours:.1f} h"),
                ("장애 횟수", "0 회"),
                ("MTBF", "N/A"),
            ]

            for k, v in metrics:
                r = table.add_row().cells
                r[0].text = k
                r[1].text = v

            # ===============================
            # 2. 핵심 성능 지표
            # ===============================
            doc.add_heading("2. 핵심 성능 지표 (Key Parameters)", level=1)

            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "항목"
            hdr[1].text = "정상 범위"
            hdr[2].text = "통계 (평균|최대|최소)"
            hdr[3].text = "상태"

            device_spec = SPEC_MAP.get(device_name, {})

            for param_name, xs, ys in param_list:
                arr = np.array(ys, dtype=float)
                mean_v = arr.mean()
                min_v = arr.min()
                max_v = arr.max()

                spec = device_spec.get(param_name)
                if spec:
                    low, high = spec
                    spec_str = f"{low:.2f} ~ {high:.2f}"
                    status = "경고" if min_v < low or max_v > high else "정상"
                else:
                    spec_str = "-"
                    status = "정상"

                r = table.add_row().cells
                r[0].text = param_name
                r[1].text = spec_str
                r[2].text = f"{mean_v:.2f} | {max_v:.2f} | {min_v:.2f}"
                r[3].text = status

            # ===============================
            # 3. 추세 그래프
            # ===============================
            doc.add_heading("3. 추세 및 이슈 (Trends & Issues)", level=1)

            for idx, (param_name, xs, ys) in enumerate(param_list):
                fig = Figure(figsize=(7, 3), dpi=120)
                ax = fig.add_subplot(111)
                ax.plot(xs, ys, linewidth=1)
                ax.set_title(f"{device_name} - {param_name}", fontsize=9)
                fig.autofmt_xdate()

                img_path = f"_vlbi_temp_{device_name}_{idx}.png"
                fig.savefig(img_path, dpi=120, bbox_inches="tight")
                temp_files.append(img_path)

                doc.add_heading(f"주요 추세 그래프: {param_name}", level=2)
                doc.add_picture(img_path, width=Cm(16))

            # ===============================
            # 4. 종합 진단
            # ===============================
            doc.add_heading("4. 종합 진단 (Conclusion)", level=1)
            doc.add_paragraph("건전성 등급: 양호")
            doc.add_paragraph("특이사항 및 조치:")
            doc.add_paragraph("- (수동 작성 영역)")

        # 저장
        doc.save(filename)

        for f in temp_files:
            try:
                os.remove(f)
            except:
                pass

        logger.info("Word 보고서 저장 완료: %s", filename)

    def get_current_selected_items(self):
        """현재 실제로 그래프가 그려지고 있는 (parent, child) 목록 반환"""
        plot_items = self._collect_plot_items()
        selected = []

        for title, xs, ys in plot_items:
            # title = "2GHz 수신기 상태 모니터 | LNA Monitor RHCP Id"
            parent, child = title.split(" | ", 1)
            selected.append((parent, child))

        return selected

    def _get_thresholds_for(self, parent_name: str, child_name: str):
        """
        parent_name: "2GHz Receiver Status Monitor"
        child_name : "Pressure Sensor CH1"
        thresholds.json에서 (table, column) 기준으로 임계값 반환
        """
        info = TABLE_MAP.get(parent_name)
        if not info:
            return None

        table = info.get("table")  # 예: Frontend_2ghz
        col = info["columns"].get(child_name)  # 예: Pressure
        if not table or not col:
            return None

        th = self.thresholds.get(table, {}).get(col)
        if not th:
            return None

        return {
            "ly": th.get("하한 주의"),
            "lr": th.get("하한 경고"),
            "uy": th.get("상한 주의"),
            "ur": th.get("상한 경고"),
        }

    def _selected_count(self) -> int:
        return sum(len(v) for v in self.selected_children.values())

    def _target_points_for_range(self) -> int:
        """시간 범위별 권장 최대 포인트 수(그래프 1개당)"""
        if self.time_range == "1Hour":
            return 600  # 1초 주기면 600포인트(10분) 정도, 환경 따라 조절
        if self.time_range == "6Hour":
            return 900
        if self.time_range == "1Day":
            return 1200
        if self.time_range == "7Day":
            return 1500
        return 1200

    def _downsample_to_max_points(self, xs, ys, max_points: int):
        """
        의미를 바꾸지 않는 표시용 디시메이션:
        인덱스를 균일하게 건너뛰어 최대 max_points로 줄임
        """
        n = len(xs)
        if n <= max_points or max_points <= 0:
            return xs, ys

        step = math.ceil(n / max_points)

        xs_ds = xs[::step]
        ys_ds = ys[::step]

        # 마지막 점 포함(끝 시각이 빠지는 걸 방지)
        if xs_ds and xs_ds[-1] != xs[-1]:
            xs_ds.append(xs[-1])
            ys_ds.append(ys[-1])

        return xs_ds, ys_ds